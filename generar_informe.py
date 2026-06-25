"""Genera el informe de la Tarea T02.03 en formato .docx siguiendo la plantilla
oficial de la UPS (secciones 1-10). Luego se convierte a PDF.

Uso:
    pip install python-docx
    python generar_informe.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------- CONFIG (datos del grupo) ----------------
CONFIG = {
    "grupo": "AKETOY",
    "docente": "Guillermo Pizarro",
    "fecha": "2026-06-24",
    "repositorio_url": "https://github.com/BarciaSchool/T02.03.git",
    "integrantes": [
        "Abatte Kelly",
        "Barcia Adrian",
        "Huambo Cesar",
        "Totoy Victor",
    ],
    "archivo_estudiante": "BarciaAdrian",
}

CONCLUSIONES = (
    "El desarrollo de la presente aplicación permitió aplicar de manera práctica los "
    "principios de la ingeniería de software estudiados en la asignatura, materializando "
    "un servicio REST funcional para la gestión de categorías. La implementación del "
    "backend bajo una arquitectura en capas —modelo, repositorio, servicio y controlador— "
    "evidenció las ventajas de la separación de responsabilidades: cada componente cumple "
    "una función específica, lo que facilita el mantenimiento, las pruebas y la "
    "escalabilidad del sistema. El uso de FastAPI como framework agilizó la construcción de "
    "los servicios y proporcionó documentación interactiva automática mediante "
    "Swagger/OpenAPI, cumpliendo con el requerimiento de publicar servicios documentados. "
    "La gestión del proyecto mediante un repositorio de código con control de versiones "
    "resultó fundamental para coordinar el trabajo del equipo; la definición y asignación de "
    "tareas, junto con los commits incrementales, permitió evidenciar la participación de "
    "todos los integrantes y mantener la trazabilidad del desarrollo. Asimismo, la "
    "contenerización con Docker garantizó un despliegue reproducible e independiente del "
    "entorno. En conjunto, la tarea consolidó competencias técnicas en el diseño e "
    "implementación de APIs y, sobre todo, en el trabajo colaborativo propio de un equipo de "
    "desarrollo de software, reforzando la importancia de las buenas prácticas de "
    "documentación y versionamiento en proyectos reales."
)

AZUL = RGBColor(0x0C, 0x15, 0xBB)


def h(doc, texto, size=13):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = AZUL
    return p


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Portada
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Tarea T02.03\nConstrucción de aplicación de software")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = AZUL
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "Universidad Politécnica Salesiana\n"
        "Carrera de Computación · Ingeniería de Software · 2026"
    ).italic = True
    doc.add_page_break()

    # 1. Fecha
    h(doc, "1. Fecha")
    doc.add_paragraph(CONFIG["fecha"])

    # 2. Grupo base
    h(doc, "2. Grupo base")
    doc.add_paragraph(f"Grupo {CONFIG['grupo']}")

    # 3. Integrantes
    h(doc, "3. Integrantes")
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Table Grid"
    tabla.rows[0].cells[0].text = "Apellidos y Nombres"
    tabla.rows[0].cells[1].text = "Calificación"
    for integ in CONFIG["integrantes"]:
        celdas = tabla.add_row().cells
        celdas[0].text = integ
        celdas[1].text = ""

    # 4. Objetivos específicos
    h(doc, "4. Objetivos específicos")
    doc.add_paragraph(
        "4.1. Crear aplicaciones utilizando herramientas de Ingeniería de software.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "4.2. Comprender el manejo de un equipo de desarrollo de software.",
        style="List Bullet",
    )

    # 5. Tareas específicas
    h(doc, "5. Tareas específicas realizadas")
    doc.add_paragraph("5.1. Se siguieron las instrucciones del ejercicio.", style="List Bullet")
    doc.add_paragraph(
        "Se creó el repositorio de código y se definieron tareas de seguimiento asignadas a "
        "cada integrante del grupo. Se implementó el backend de la aplicación de gestión de "
        "categorías con arquitectura en capas (modelo → repositorio → servicio → controlador), "
        "se publicaron los servicios REST con documentación Swagger/OpenAPI y se desplegó la "
        "solución mediante contenedores Docker."
    )

    # 6. Entregables
    h(doc, "6. Entregables")
    doc.add_paragraph(
        f"6.1. Archivo en PDF: T02_03_Grupo{CONFIG['grupo']}_{CONFIG['archivo_estudiante']}.pdf",
        style="List Bullet",
    )
    doc.add_paragraph(
        "6.2. Aplicación de software almacenada en repositorio de código.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "6.3. La carga a la plataforma AVAC es individual aunque la tarea sea grupal.",
        style="List Bullet",
    )

    # 7. Instrucciones / cumplimiento
    h(doc, "7. Instrucciones (cumplimiento)")
    instrucciones = [
        "7.1. Se creó un repositorio de código y se definieron tareas de seguimiento "
        "asignadas a cada integrante del grupo (ver docs/tareas.md).",
        "7.2. Se desarrolló un software que cumple con los requerimientos de la Tarea 02.01 "
        "y las descripciones de diseño de la Tarea 02.02.",
        "7.3. El desarrollo consistió en la implementación del backend con el framework "
        "FastAPI (opción 7.3.2 del enunciado).",
        "7.4. Se publicaron los servicios con documentación mediante Swagger/OpenAPI "
        "(rutas /docs y /redoc).",
        "7.5. Se utilizó un esquema modelo/repositorio/servicio/controlador similar al "
        "definido en la Figura 1 del enunciado.",
    ]
    for ins in instrucciones:
        doc.add_paragraph(ins, style="List Bullet")

    # 8. Código fuente
    h(doc, "8. Código fuente")
    p = doc.add_paragraph("Repositorio: ")
    p.add_run(CONFIG["repositorio_url"]).bold = True

    # 9. Conclusiones
    h(doc, "9. Conclusiones")
    doc.add_paragraph(CONCLUSIONES)

    # 10. Rúbrica
    h(doc, "10. Rúbrica de calificación")
    rub = doc.add_table(rows=1, cols=6)
    rub.style = "Table Grid"
    enc = ["Criterios", "No Presenta", "Nivel Bajo", "Nivel Medio", "Nivel Alto", "Total"]
    for i, e in enumerate(enc):
        rub.rows[0].cells[i].text = e
    filas = [
        ["Repositorio correctamente documentado con los requerimientos, "
         "diseño y tareas definidas por usuario.", "0", "0.5", "1", "2", ""],
        ["Al menos 20 commits en el repositorio que aseguren que han "
         "trabajado todos los compañeros.", "0", "0.2", "0.5", "1", ""],
        ["Servicios funcionales y correctamente desplegados.", "0", "0.2", "0.5", "1", ""],
    ]
    for fila in filas:
        celdas = rub.add_row().cells
        for i, v in enumerate(fila):
            celdas[i].text = v

    nombre = f"T02_03_Grupo{CONFIG['grupo']}_{CONFIG['archivo_estudiante']}.docx"
    doc.save(nombre)
    print(f"Generado: {nombre}")


if __name__ == "__main__":
    main()
