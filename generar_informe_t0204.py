"""Genera los informes de la Tarea T02.04 (pruebas unitarias) rellenando la
PLANTILLA OFICIAL de la UPS (portada y logo incluidos). Produce un .docx por
integrante; luego se convierten a PDF (ver bloque PowerShell en el README de uso).

Uso:
    pip install python-docx
    python generar_informe_t0204.py
"""
from docx import Document
from docx.oxml.ns import qn

# ---------------- CONFIG (datos del grupo) ----------------
PLANTILLA = r"T02_04_GrupoXX_Apellido1Nombre1.docx"

CONFIG = {
    "grupo": "AKETOY",
    "docente": "Guillermo Pizarro",
    "fecha": "2026-06-29",
    "repositorio_url": "https://github.com/BarciaSchool/T02.03.git",
    "cobertura_lograda": "97.84%",
    "integrantes": [
        "Abatte Kelly",
        "Barcia Adrian",
        "Huambo Cesar",
        "Totoy Victor",
    ],
}

CONCLUSIONES = (
    "La implementación de pruebas unitarias sobre la API de categorías permitió "
    "verificar de forma automática y reproducible el correcto funcionamiento de cada "
    "capa del sistema. Se diseñaron pruebas independientes para el repositorio, el "
    "servicio y el controlador, aplicando el principio de aislamiento mediante objetos "
    "simulados (mocks) que reemplazan las dependencias reales, de modo que la lógica de "
    "negocio del servicio se valida sin depender de la base de datos. Para ello se "
    "emplearon Pytest como ejecutor principal, unittest y unittest.mock para el estilo "
    "basado en clases y la simulación de dependencias, doctest para validar utilidades a "
    "partir de ejemplos en la documentación, y Coverage.py para medir el porcentaje de "
    "código ejercitado por las pruebas. El análisis de cobertura superó el umbral "
    "exigido del 60% de los métodos, evidenciando que las rutas principales y los casos "
    "de error, como la solicitud de recursos inexistentes, están cubiertos. Este proceso "
    "demostró el valor de las pruebas automatizadas como red de seguridad ante futuros "
    "cambios: permiten detectar regresiones de manera temprana, documentan el "
    "comportamiento esperado del sistema y aumentan la confianza en el despliegue. "
    "Asimismo, el trabajo colaborativo en la definición de las pruebas, reflejado en el "
    "historial de commits, reforzó las prácticas de calidad propias de un equipo de "
    "desarrollo de software."
)


def replace_in_paragraph(p, old, new):
    """Reemplaza `old` por `new` en un párrafo, preservando el formato.
    Funciona aunque `old` esté repartido en varios runs."""
    if old in p.text:
        for run in p.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                return True
        if p.runs:
            p.runs[0].text = p.text.replace(old, new)
            for run in p.runs[1:]:
                run.text = ""
            return True
    return False


def build(doc_template_path, archivo_estudiante):
    doc = Document(doc_template_path)

    for p in doc.paragraphs:
        # 1. Fecha
        if p.text.strip() == "Fecha:":
            p.add_run(f" {CONFIG['fecha']}")
        # 2. Grupo base
        replace_in_paragraph(p, "Grupo XX", f"Grupo {CONFIG['grupo']}")
        # Nombre del archivo PDF
        replace_in_paragraph(p, "GrupoXX", f"Grupo{CONFIG['grupo']}")
        replace_in_paragraph(p, "Apellido1Nombre1", archivo_estudiante)
        # Enlace del repositorio (sustituye la línea de EJEMPLO).
        # El texto de ejemplo incluye un hipervínculo (w:hyperlink), por eso se
        # eliminan tanto los runs como los hipervínculos antes de escribir el real.
        if p.text.strip().startswith("EJEMPLO:"):
            for child in list(p._p):
                if child.tag in (qn("w:r"), qn("w:hyperlink")):
                    p._p.remove(child)
            p.add_run(f"Repositorio: {CONFIG['repositorio_url']}")

    # Integrantes (primera tabla, columna "Apellidos y Nombres")
    tabla = doc.tables[0]
    filas_datos = tabla.rows[1:]
    for i, integ in enumerate(CONFIG["integrantes"]):
        if i < len(filas_datos):
            filas_datos[i].cells[0].text = integ

    # Conclusiones (insertar el texto justo después del encabezado)
    for p in doc.paragraphs:
        if p.text.strip().startswith("Conclusiones de la Tarea"):
            nota = doc.add_paragraph(
                f"Cobertura de código alcanzada: {CONFIG['cobertura_lograda']} "
                f"(umbral exigido: 60%). 41 pruebas ejecutadas con éxito."
            )
            nuevo_p = doc.add_paragraph(CONCLUSIONES)
            p._p.addnext(nuevo_p._p)
            nuevo_p._p.addprevious(nota._p)
            break

    nombre = f"T02_04_Grupo{CONFIG['grupo']}_{archivo_estudiante}.docx"
    doc.save(nombre)
    print(f"Generado: {nombre}")
    return nombre


def main():
    # Un informe por integrante (la carga al AVAC es individual).
    for integ in CONFIG["integrantes"]:
        archivo_estudiante = integ.replace(" ", "")
        build(PLANTILLA, archivo_estudiante)


if __name__ == "__main__":
    main()
